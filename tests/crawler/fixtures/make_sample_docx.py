# tests/crawler/fixtures/make_sample_docx.py
from pathlib import Path
try:
    import docx  # python-docx
except Exception:  # sem extra, skip silencioso
    print("python-docx not installed; skipping docx creation.")
    raise SystemExit(0)

p = Path(__file__).parent / "sample.docx"
d = docx.Document()
d.add_heading("Sample DOCX", 0)
d.add_paragraph("Hello Aurora Crawler.")
d.save(p)
print(f"Wrote {p}")
